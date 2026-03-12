import json
import re
from pathlib import Path

import numpy as np
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

from cmj_framework.utils.app_logging import make_session_log_file, append_to_path
from cmj_framework.utils.runtime_paths import (
    config_file,
    export_resource_file,
    export_user_resources_dir,
    ensure_dir,
)


# =========================================================
# Exportable parameter definitions
# =========================================================

PERF_METRICS = [
    ("Sprunghöhe (cm)", ["jump_height"]),
    ("Körpergewicht (N)", ["weight"]),
    ("Absprungsgeschwindigkeit (m/s)", ["Takeoff_velocity"]),
    ("RSI Modified", ["RSI_modified"]),
    ("Zeit bis zum Absprung (s)", ["PhaseDurations", "Time_to_takeoff"]),
    ("Dauer exzentrisch (s)", ["PhaseDurations", "Eccentric_duration"]),
    ("Dauer konzentrisch (s)", ["PhaseDurations", "Concentric_duration"]),
    ("RFD max (N/s)", ["RFD", "RFD_max"]),
    ("RFD 0–100ms (N/s)", ["RFD", "RFD_0_100ms"]),
    ("RFD 0–200ms (N/s)", ["RFD", "RFD_0_200ms"]),
    ("Mittlere Leistung (W)", ["ConcentricPower", "Power_mean"]),
    ("Maximale Leistung (W)", ["ConcentricPower", "Power_peak"]),
    ("Maximale Belastungsrate (N/s)", ["LandingMetrics", "LoadingRate_peak"]),
    ("Maximale Landekraft (N)", ["LandingMetrics", "PeakLandingForce"]),
]

PEAKFORCE_METRICS = [
    ("Maximale Bremskraft (N)", ["PeakForce", "braking"]),
    ("Maximale Verzögerungskraft (N)", ["PeakForce", "deceleration"]),
    ("Maximale Antriebskraft (N)", ["PeakForce", "propulsion"]),
    ("Maximale Landekraft (N)", ["PeakForce", "landing"]),
]


def get_all_export_parameter_labels() -> list[str]:
    """Return all parameter labels that can appear in the performance table."""
    return [label for label, _ in (PERF_METRICS + PEAKFORCE_METRICS)]


# =========================================================
# Paths / resources / config / logging
# =========================================================

def load_cfg() -> dict:
    """Load the export JSON configuration file."""
    cfg_path = config_file("word_report_config.json")
    try:
        with open(cfg_path, "r", encoding="utf-8") as file:
            return json.load(file)
    except Exception as exc:
        raise RuntimeError(f"Fehler beim Laden der Word-Export-Konfiguration: {exc}")


CFG = load_cfg()
LOG_FILE = make_session_log_file("word_export")


def log(msg: str) -> None:
    """Append a message to the Word export session log file."""
    append_to_path(LOG_FILE, msg)


def get_writable_export_resource_dir() -> Path:
    """
    Return the writable directory for user-editable export resources.

    This folder is persistent and independent from the bundle.
    """
    return ensure_dir(export_user_resources_dir())


def ensure_editable_markdown_file(filename: str) -> str:
    """
    Ensure a writable copy of a bundled markdown file exists and return its path.
    """
    writable_dir = get_writable_export_resource_dir()
    dst = writable_dir / filename

    if dst.exists():
        return str(dst)

    src = export_resource_file(filename)
    if src.exists():
        dst.write_text(src.read_text(encoding="utf-8"), encoding="utf-8")
    else:
        dst.write_text("", encoding="utf-8")

    return str(dst)


# =========================================================
# Generic helpers
# =========================================================

def infer_patient_from_path(json_path: str) -> str:
    """
    Infer patient name from canonical CMJ folder structure:
    base/patient/session/processed/file.json
    """
    try:
        return Path(json_path).parents[2].name or "Unknown"
    except Exception:
        return "Unknown"


def safe_float(value, default=None):
    """Convert to float safely; return default on failure."""
    try:
        if value is None:
            return default
        return float(value)
    except Exception:
        return default


def fetch(data, path, default=None):
    """Fetch nested dict value using a list of keys."""
    current = data
    for key in path:
        if not isinstance(current, dict) or key not in current:
            return default
        current = current[key]
    return current


def sanitize_filename(name: str) -> str:
    """Replace invalid filename characters."""
    for bad in ['\\', '/', ':', '*', '?', '"', '<', '>', '|']:
        name = name.replace(bad, "_")
    return name.strip()


# =========================================================
# Word layout helpers
# =========================================================

def format_A4(section) -> None:
    """Apply A4 page dimensions based on current orientation."""
    width = Mm(CFG["document"]["page"]["width_mm"])
    height = Mm(CFG["document"]["page"]["height_mm"])

    if section.orientation == WD_ORIENT.PORTRAIT:
        section.page_width = width
        section.page_height = height
    else:
        section.page_width = height
        section.page_height = width


def remove_table_borders(table) -> None:
    """Hard-remove all borders for a python-docx table."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    try:
        table.style = "Table Normal"
    except Exception:
        pass

    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = tbl_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tbl_borders.append(element)

        element.set(qn("w:val"), "nil")
        element.set(qn("w:sz"), "0")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "auto")


def remove_cell_borders(cell) -> None:
    """Hard-remove borders for a single cell."""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()

    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("top", "left", "bottom", "right"):
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)

        element.set(qn("w:val"), "nil")
        element.set(qn("w:sz"), "0")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "auto")


def header(section, left_text: str, right_text: str) -> None:
    """
    Clean header: left/right aligned text using a 2-column borderless table.
    """
    section.header.is_linked_to_previous = False
    hdr = section.header

    for paragraph in hdr.paragraphs:
        try:
            paragraph.clear()
        except Exception:
            pass

    usable_width = section.page_width - section.left_margin - section.right_margin
    table = hdr.add_table(rows=1, cols=2, width=usable_width)
    table.autofit = True

    try:
        table.style = "Table Normal"
    except Exception:
        pass

    cell_left = table.cell(0, 0)
    cell_left.text = left_text
    cell_left.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    cell_right = table.cell(0, 1)
    cell_right.text = right_text
    cell_right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    remove_table_borders(table)
    remove_cell_borders(cell_left)
    remove_cell_borders(cell_right)


def add_patient_date_row(doc, section, patient: str, date_str: str) -> None:
    """Add patient/date row at top of the page."""
    del section

    table = doc.add_table(rows=1, cols=2)
    table.autofit = True

    table.cell(0, 0).text = patient
    table.cell(0, 1).text = date_str

    table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    remove_table_borders(table)
    remove_cell_borders(table.cell(0, 0))
    remove_cell_borders(table.cell(0, 1))

    doc.add_paragraph("")


def cell_border(cell, **kwargs) -> None:
    """Apply border settings to a cell."""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()

    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("top", "bottom", "left", "right"):
        if edge not in kwargs:
            continue

        edge_data = kwargs[edge]
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)

        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))


def _rgb_to_hex(rgb) -> str:
    """Convert RGBColor or hex-like input to uppercase 6-digit hex."""
    if rgb is None:
        return "FFFFFF"

    if isinstance(rgb, RGBColor):
        return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"

    s = str(rgb).strip().lstrip("#")
    if len(s) == 6 and all(c in "0123456789abcdefABCDEF" for c in s):
        return s.upper()

    return "FFFFFF"


def set_cell_background(cell, rgb) -> None:
    """Set cell background fill."""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()

    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)

    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _rgb_to_hex(rgb))


def set_cell_vertical_center(cell) -> None:
    """Vertically center cell content."""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    v_align = OxmlElement("w:vAlign")
    v_align.set(qn("w:val"), "center")
    tc_pr.append(v_align)


def compact_cell(cell, font_size: int) -> None:
    """Apply compact paragraph spacing and font size to cell content."""
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Mm(1)
        paragraph.paragraph_format.space_after = Mm(1)
        paragraph.paragraph_format.line_spacing = 1
        for run in paragraph.runs:
            run.font.size = Pt(font_size)


def set_row_height(row, height_mm: float) -> None:
    """Set row height in millimeters."""
    tr = row._tr
    tr_pr = tr.get_or_add_trPr()
    tr_height = OxmlElement("w:trHeight")
    tr_height.set(qn("w:val"), str(int(height_mm * 56.7)))
    tr_height.set(qn("w:hRule"), "atLeast")
    tr_pr.append(tr_height)


def set_cell_width(cell, width_cm: float) -> None:
    """Set cell width in centimeters."""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_pr.append(tc_w)


def create_table1_look_compact(doc, df):
    """Create the asymmetry table with merged sections and compact styling."""
    rows, cols = df.shape
    table = doc.add_table(rows=rows + 1, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    table.style = "Table Grid"

    header_color = RGBColor(217, 217, 217)
    row_colors = [RGBColor(255, 255, 255), RGBColor(245, 245, 245)]
    last_col = cols - 1
    border_style = {"sz": 4, "val": "single", "color": "000000"}

    for j, col in enumerate(df.columns):
        cell = table.cell(0, j)
        cell.text = str(col)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if paragraph.runs:
            paragraph.runs[0].font.bold = True
            paragraph.runs[0].font.size = Pt(9)

        set_cell_background(cell, header_color)
        cell_border(cell, top=border_style, bottom=border_style, left=border_style, right=border_style)
        set_row_height(table.rows[0], 10)

    i = 0
    block_idx = 0
    while i < rows:
        start_i = i
        group_val = df.iloc[i, 0]
        row_color = row_colors[block_idx % 2]

        while i < rows and df.iloc[i, 0] == group_val:
            i += 1
        end_i = i - 1

        merged_cell0 = table.cell(start_i + 1, 0).merge(table.cell(end_i + 1, 0))
        merged_cell0.text = str(group_val)
        set_cell_background(merged_cell0, row_color)
        cell_border(merged_cell0, top=border_style, bottom=border_style, left=border_style, right=border_style)
        merged_cell0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if merged_cell0.paragraphs[0].runs:
            merged_cell0.paragraphs[0].runs[0].font.size = Pt(9)

        r = start_i
        while r <= end_i:
            if r + 1 <= end_i:
                phase_cell = table.cell(r + 1, 1).merge(table.cell(r + 2, 1))
                phase_cell.text = str(df.iloc[r, 1])
                set_cell_background(phase_cell, row_color)
                cell_border(phase_cell, top=border_style, bottom=border_style, left=border_style, right=border_style)
                phase_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if phase_cell.paragraphs[0].runs:
                    phase_cell.paragraphs[0].runs[0].font.size = Pt(9)

                asym_cell = table.cell(r + 1, last_col).merge(table.cell(r + 2, last_col))
                asym_cell.text = str(df.iloc[r, last_col])
                set_cell_background(asym_cell, row_color)
                cell_border(asym_cell, top=border_style, bottom=border_style, left=border_style, right=border_style)
                asym_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if asym_cell.paragraphs[0].runs:
                    asym_cell.paragraphs[0].runs[0].font.size = Pt(8)

            r += 2

        for r in range(start_i, end_i + 1):
            for c in range(2, last_col):
                cell = table.cell(r + 1, c)
                cell.text = str(df.iloc[r, c])
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_cell_background(cell, row_color)
                cell_border(cell, top=border_style, bottom=border_style, left=border_style, right=border_style)
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].font.size = Pt(9)
                set_row_height(table.rows[r + 1], 11.5)

        block_idx += 1

    for r in range(rows + 1):
        set_cell_width(table.cell(r, 0), 2.2)
        set_cell_width(table.cell(r, 1), 1.8)
        set_cell_width(table.cell(r, 2), 1.0)

        for j in range(3, last_col):
            set_cell_width(table.cell(r, j), 1.15)

        set_cell_width(table.cell(r, last_col), 2.4)

    return table


def create_classic_table_with_stats(doc, df, add_stats: bool = True):
    """Create the performance table with optional mean ± std summary column."""
    rows, cols = df.shape
    table = doc.add_table(rows=rows + 1, cols=cols + (1 if add_stats else 0))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    row_colors = [RGBColor(245, 245, 245), RGBColor(255, 255, 255)]

    for j, col in enumerate(df.columns):
        cell = table.cell(0, j)
        cell.text = str(col)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].font.bold = True

        set_cell_vertical_center(cell)
        cell_border(
            cell,
            top={"sz": 12, "val": "single", "color": "000000"},
            bottom={"sz": 12, "val": "single", "color": "000000"},
            left={"sz": 12, "val": "single", "color": "000000"},
            right={"sz": 12, "val": "single", "color": "000000"},
        )
        set_cell_background(cell, RGBColor(217, 217, 217))

    if add_stats:
        cell = table.cell(0, cols)
        cell.text = "x̄ ± s"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].font.bold = True

        set_cell_vertical_center(cell)
        cell_border(
            cell,
            top={"sz": 12, "val": "single", "color": "000000"},
            bottom={"sz": 12, "val": "single", "color": "000000"},
            left={"sz": 12, "val": "single", "color": "000000"},
            right={"sz": 12, "val": "single", "color": "000000"},
        )
        set_cell_background(cell, RGBColor(217, 217, 217))

    for r in range(rows):
        row_color = row_colors[r % 2]
        for c in range(cols):
            cell = table.cell(r + 1, c)
            value = df.iloc[r, c]
            cell.text = str(value)
            set_cell_vertical_center(cell)
            set_cell_background(cell, row_color)
            compact_cell(cell, font_size=9)
            cell_border(
                cell,
                top={"sz": 6, "val": "single", "color": "000000"},
                bottom={"sz": 6, "val": "single", "color": "000000"},
                left={"sz": 6, "val": "single", "color": "000000"},
                right={"sz": 6, "val": "single", "color": "000000"},
            )

        if add_stats:
            try:
                numeric_vals = []
                for c in range(1, cols):
                    value = safe_float(df.iloc[r, c], default=None)
                    if value is not None:
                        numeric_vals.append(value)

                cell = table.cell(r + 1, cols)
                if numeric_vals:
                    mean = round(float(np.mean(numeric_vals)), 2)
                    std = round(float(np.std(numeric_vals)), 2)
                    cell.text = f"{mean} ± {std}"
                else:
                    cell.text = "_"

                set_cell_vertical_center(cell)
                set_cell_background(cell, row_color)
                compact_cell(cell, font_size=8)
                cell_border(
                    cell,
                    top={"sz": 6, "val": "single", "color": "000000"},
                    bottom={"sz": 6, "val": "single", "color": "000000"},
                    left={"sz": 6, "val": "single", "color": "000000"},
                    right={"sz": 6, "val": "single", "color": "000000"},
                )
            except Exception:
                pass

        set_row_height(table.rows[r + 1], 10.5)

    for r in range(rows + 1):
        set_cell_width(table.cell(r, 0), 5.0)

        for c in range(1, cols):
            set_cell_width(table.cell(r, c), 1.35)

        if add_stats:
            set_cell_width(table.cell(r, cols), 2.4)

    return table


# =========================================================
# Markdown insertion
# =========================================================

def insert_markdown(doc, md_path: str) -> None:
    """Insert simplified markdown content into the Word document."""
    with open(md_path, "r", encoding="utf-8") as file:
        lines = file.readlines()

    list_buffer: list[str] = []

    def add_inline_runs(paragraph, text: str) -> None:
        tokens = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
        for token in tokens:
            if not token:
                continue

            if token.startswith("**") and token.endswith("**") and len(token) >= 4:
                run = paragraph.add_run(token[2:-2])
                run.bold = True
            elif token.startswith("*") and token.endswith("*") and len(token) >= 2:
                run = paragraph.add_run(token[1:-1])
                run.italic = True
            else:
                paragraph.add_run(token)

    def style_body_paragraph(paragraph) -> None:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.line_spacing = 1.25
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.space_before = Pt(0)

    def style_heading(paragraph, level: int) -> None:
        paragraph.paragraph_format.space_before = Pt(10 if level == 1 else 8)
        paragraph.paragraph_format.space_after = Pt(4)

    def flush_list() -> None:
        nonlocal list_buffer
        for item in list_buffer:
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing = 1.15
            add_inline_runs(paragraph, item)
        list_buffer = []

    for raw_line in lines:
        line = raw_line.rstrip("\n")
        stripped = line.strip()

        if not stripped:
            flush_list()
            continue

        if stripped == "---":
            flush_list()
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.add_run().add_break()
            continue

        if stripped.startswith("### "):
            flush_list()
            paragraph = doc.add_heading(level=3)
            add_inline_runs(paragraph, stripped[4:])
            style_heading(paragraph, 3)
            continue

        if stripped.startswith("## "):
            flush_list()
            paragraph = doc.add_heading(level=2)
            add_inline_runs(paragraph, stripped[3:])
            style_heading(paragraph, 2)
            continue

        if stripped.startswith("# "):
            flush_list()
            paragraph = doc.add_heading(level=1)
            add_inline_runs(paragraph, stripped[2:])
            style_heading(paragraph, 1)
            continue

        if stripped.startswith("- "):
            list_buffer.append(stripped[2:])
            continue

        flush_list()
        paragraph = doc.add_paragraph()
        style_body_paragraph(paragraph)
        add_inline_runs(paragraph, stripped)

    flush_list()
