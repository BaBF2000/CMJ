# -*- coding: utf-8 -*-
import argparse
import json
import os
import platform
import traceback
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Cm, Inches, Pt

from src.cmj_framework.export.word_report_helpers import (
    CFG,
    LOG_FILE,
    log,
    ensure_editable_markdown_file,
    infer_patient_from_path,
    safe_float,
    fetch,
    sanitize_filename,
    format_A4,
    header,
    add_patient_date_row,
    insert_markdown,
    create_classic_table_with_stats,
    create_table1_look_compact,
    PERF_METRICS,
    PEAKFORCE_METRICS,
)
from src.cmj_framework.utils.runtime_paths import export_resource_file


log("----- cmj_word_report gestartet -----")
log(f"CWD        : {os.getcwd()}")
log(f"LOG_FILE   : {LOG_FILE}")

def Report(
    json_path: str,
    patient_override: str | None = None,
    enabled_parameters: list[str] | None = None,
    parameter_md_path: str | None = None,
    phases_md_path: str | None = None,
):
    """
    Generate the Word report from a *_combined.json file.
    """
    log("Report() gestartet")
    log(f"JSON path: {json_path}")

    with open(json_path, "r", encoding="utf-8") as file:
        data = json.load(file)

    metrics_by_trial = data.get("metrics_by_trial", {})
    if not metrics_by_trial:
        raise RuntimeError("Ungültige JSON: 'metrics_by_trial' fehlt oder ist leer.")

    trials = list(metrics_by_trial.keys())
    date_str = data.get("session_date") or ""
    patient = patient_override or data.get("patient_name") or infer_patient_from_path(json_path)

    phases_md_path = phases_md_path or ensure_editable_markdown_file("phases.md")
    parameter_md_path = parameter_md_path or ensure_editable_markdown_file("parameter.md")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = CFG["document"]["font_name"]
    style.font.size = Pt(CFG["document"]["font_size"])

    section = doc.sections[0]
    format_A4(section)

    margins = CFG["document"]["page"]["margins_cm"]
    section.top_margin = Cm(margins["top"])
    section.bottom_margin = Cm(margins["bottom"])
    section.left_margin = Cm(margins["left"])
    section.right_margin = Cm(margins["right"])

    usable_width = (
        section.page_width.inches
        - section.left_margin.inches
        - section.right_margin.inches
    )

    section.different_first_page_header_footer = True
    first_header = section.first_page_header
    paragraph = first_header.paragraphs[0]
    run = paragraph.add_run()

    logo_path = export_resource_file("cmj_banner_placeholder.png")
    if not logo_path.exists():
        raise FileNotFoundError(f"Logo nicht gefunden: {logo_path}")

    run.add_picture(str(logo_path), width=Inches(usable_width))
    log(f"Headerbild geladen: {logo_path}")

    header(section, left_text="", right_text="")

    doc.add_heading("Counter Movement Jump Bericht", level=0)
    add_patient_date_row(doc, section, patient, date_str)

    graph_path = export_resource_file("munster_graph.png")
    if not graph_path.exists():
        raise FileNotFoundError(f"Graph nicht gefunden: {graph_path}")

    doc.add_picture(str(graph_path), width=Inches(usable_width))
    log(f"Graph geladen: {graph_path}")

    if os.path.exists(phases_md_path):
        insert_markdown(doc, phases_md_path)
        log(f"Markdown der Phasen eingefügt: {phases_md_path}")
    else:
        log("phases.md fehlt (übersprungen)")

    section1 = doc.add_section()
    section1.different_first_page_header_footer = False
    header(section1, left_text=patient, right_text=date_str)

    rows = []
    asym_metrics_map = {"Impulse": "Impuls", "Power": "Leistung"}
    phase_map = {"Takeoff": "Absprung", "Landing": "Landung"}

    for group_en, group_de in asym_metrics_map.items():
        for phase_en, phase_de in phase_map.items():
            left_values, right_values, deltas = [], [], []

            for trial in trials:
                trial_data = metrics_by_trial.get(trial, {})
                phase_data = fetch(trial_data, ["Asymmetry", group_en, phase_en], default=None)

                if not isinstance(phase_data, dict):
                    left_values.append("_")
                    right_values.append("_")
                    continue

                left = safe_float(phase_data.get("contribution_L_percent"), default=None)
                right = safe_float(phase_data.get("contribution_R_percent"), default=None)
                delta = safe_float(phase_data.get("delta_percent"), default=None)

                left_values.append(round(left, 2) if left is not None else "_")
                right_values.append(round(right, 2) if right is not None else "_")

                if delta is not None:
                    deltas.append(delta)

            mean = round(float(np.mean(deltas)), 2) if deltas else "_"
            std = round(float(np.std(deltas)), 2) if deltas else "_"
            asym_text = f"{mean} ± {std}" if deltas else "_"

            rows.append([group_de, phase_de, "L"] + left_values + [asym_text])
            rows.append([group_de, phase_de, "R"] + right_values + [""])

    columns = ["Messwert", "Phase", "Fuß"] + trials + ["Asymmetrie"]
    df_asym = pd.DataFrame(rows, columns=columns)

    perf_rows = []
    selected_labels = set(enabled_parameters) if enabled_parameters else None

    for label, path in PERF_METRICS + PEAKFORCE_METRICS:
        if selected_labels is not None and label not in selected_labels:
            continue

        row = [label]
        for trial in trials:
            trial_data = metrics_by_trial.get(trial, {})
            value = fetch(trial_data, path, default=None)
            value = safe_float(value, default=None)
            row.append(round(value, 2) if value is not None else "_")
        perf_rows.append(row)

    if not perf_rows:
        raise RuntimeError("Keine Performance-Parameter ausgewählt. Bericht kann nicht erstellt werden.")

    df_perf = pd.DataFrame(perf_rows, columns=["Metrik"] + trials)

    section_landscape = doc.add_section()
    section_landscape.orientation = WD_ORIENT.LANDSCAPE
    format_A4(section_landscape)
    section_landscape.different_first_page_header_footer = False
    header(section_landscape, left_text=patient, right_text=date_str)

    create_classic_table_with_stats(doc, df_perf, add_stats=True)
    doc.add_page_break()
    create_table1_look_compact(doc, df_asym)

    section_parameter = doc.add_section()
    section_parameter.orientation = WD_ORIENT.PORTRAIT
    format_A4(section_parameter)
    section_parameter.different_first_page_header_footer = False
    header(section_parameter, left_text=patient, right_text=date_str)

    if os.path.exists(parameter_md_path):
        insert_markdown(doc, parameter_md_path)
        log(f"Markdown der Parameter eingefügt: {parameter_md_path}")
    else:
        log("parameter.md fehlt (übersprungen)")

    json_path_obj = Path(json_path)
    try:
        session_dir = json_path_obj.parent.parent
        output_dir = session_dir / "reports"
    except Exception:
        output_dir = Path(os.path.dirname(json_path)) / "reports"

    output_dir.mkdir(parents=True, exist_ok=True)

    safe_patient = sanitize_filename(patient or "Unknown_Patient")
    safe_date = sanitize_filename(data.get("session_date") or "xx.xx.xxxx")
    filename = f"{safe_patient}_{safe_date}_CMJ_Bericht.docx"
    output_path = output_dir / filename

    doc.save(str(output_path))
    log(f"Bericht erfolgreich erstellt: {output_path}")
    return str(output_path)


if __name__ == "__main__":
    try:
        parser = argparse.ArgumentParser()
        parser.add_argument("json_path", help="Pfad zur Datei *_combined.json")
        parser.add_argument("--patient", default=None, help="Patientenname überschreiben")
        args = parser.parse_args()

        json_path = args.json_path

        if not os.path.exists(json_path):
            raise FileNotFoundError(f"JSON-Datei nicht gefunden: {json_path}")

        report_path = Report(json_path, patient_override=args.patient)

        if platform.system() == "Windows":
            os.startfile(report_path)

    except Exception:
        log("!!! FEHLER !!!")
        log(traceback.format_exc())
        raise