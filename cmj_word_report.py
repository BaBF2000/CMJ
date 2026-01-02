import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches, Cm, Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm
import ast


# ------------------- Page format -------------------
def set_section_A4(section):
    """
    Forces A4 format for a given document section.
    Works for both portrait and landscape orientations.
    """
    A4_WIDTH = Mm(210)   # 21.0 cm
    A4_HEIGHT = Mm(297)  # 29.7 cm

    if section.orientation == WD_ORIENT.PORTRAIT:
        section.page_width = A4_WIDTH
        section.page_height = A4_HEIGHT
    else:  # LANDSCAPE
        section.page_width = A4_HEIGHT
        section.page_height = A4_WIDTH


# ------------------- Statistical calculations -------------------
def stat_impuls(L, R):
    """
    Computes the percentage distribution between left and right sides
    and their corresponding differences.

    Parameters:
    L = left leg values
    R = right leg values

    Returns:
    - left percentages
    - right percentages
    - absolute percentage difference per jump
    - mean of the differences
    - standard deviation of the differences
    """
    prozent_L = L * 100 / (L + R)
    prozent_R = R * 100 / (L + R)
    diff_prozent = abs(prozent_L - prozent_R).round()
    mittel_diff = np.mean(diff_prozent).round()
    std_diff = np.std(diff_prozent).round()
    return prozent_L, prozent_R, diff_prozent, mittel_diff, std_diff


# ------------------- Set table cell borders -------------------
def set_cell_border(cell, **kwargs):
    """
    Adds custom borders to a table cell.
    The user can specify borders such as top, bottom, left, and right.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            attrs = kwargs.get(edge)
            element = OxmlElement(f"w:{edge}")
            for key, val in attrs.items():
                element.set(qn(f"w:{key}"), str(val))
            tcPr.append(element)


# ------------------- Set table cell background color -------------------
def set_cell_background(cell, rgb_color: RGBColor):
    """
    Sets the background color of a table cell using XML manipulation.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # "shd" = shading element
    shd = OxmlElement('w:shd')
    # Set color in hex format (RGBColor already provides hex representation)
    shd.set(qn('w:fill'), f"{rgb_color}")
    tcPr.append(shd)


# ------------------- Create a formatted table -------------------
def create_responsive_styled_table(doc, pivot_df):
    """
    Creates a professionally formatted table with:
    - dynamic size based on the DataFrame
    - alternating row colors
    - highlighted header row
    - thin gray border lines
    """

    rows, cols_n = pivot_df.shape
    # +1 for header row, +1 for additional index column
    table = doc.add_table(rows=rows + 1, cols=cols_n + 1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True  # automatic column width adjustment

    # Color definitions
    header_color = RGBColor(217, 217, 217)
    row_colors = [RGBColor(242, 242, 242), RGBColor(255, 255, 255)]

    # --- Header row ---
    hdr = table.rows[0]
    hdr.cells[0].text = "Measurement"
    hdr.cells[0].paragraphs[0].runs[0].font.bold = True
    hdr.cells[0].paragraphs[0].alignment = 0
    set_cell_background(hdr.cells[0], header_color)

    # Column headers from DataFrame
    for j, col_name in enumerate(pivot_df.columns):
        cell = hdr.cells[j + 1]
        cell.text = str(col_name)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = 1
        set_cell_background(cell, header_color)

    # --- Table body ---
    for i, idx in enumerate(pivot_df.index):
        row_color = row_colors[i % 2]  # alternating row colors

        # First column = row label
        cell = table.cell(i + 1, 0)
        cell.text = idx
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = 0
        set_cell_background(cell, row_color)

        # Values in remaining columns
        for j in range(cols_n):
            c = table.cell(i + 1, j + 1)
            c.text = str(pivot_df.iloc[i, j])
            c.paragraphs[0].alignment = 1
            set_cell_background(c, row_color)

    # --- Border lines ---
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"sz": 2, "val": "single", "color": "D9D9D9"},
                bottom={"sz": 2, "val": "single", "color": "D9D9D9"},
                left={"sz": 2, "val": "single", "color": "D9D9D9"},
                right={"sz": 2, "val": "single", "color": "D9D9D9"},
            )

    return table


# ------------------- Main report generation routine -------------------
def Report(dateipfad):
    """
    Generates a complete Word report for a Counter Movement Jump.
    Includes:
    - logo on the first page
    - patient information
    - statistical calculations
    - automatically generated table
    - landscape page for improved table readability
    """
    doc = Document()

    # Set global font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Set page margins
    section = doc.sections[0]
    set_section_A4(section)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(1.91)
    section.right_margin = Cm(1.91)

    # Calculate usable width for the logo
    usable_width = section.page_width.inches - section.left_margin.inches - section.right_margin.inches

    # --- First page header (with image) ---
    section.different_first_page_header_footer = True
    first_header = section.first_page_header
    par = first_header.paragraphs[0]
    # run = par.add_run()
    # run.add_picture(
    #     r"X:\Bewegungsanalysen\Praktikanten\Beigoll\Kopf Labor.jpg",
    #     width=Inches(usable_width)
    # )

    # Clear default header
    section.header.paragraphs[0].text = ""

    # --- Report title ---
    doc.add_heading("Counter Movement Jump Report", level=0)

    # Load CSV
    df = pd.read_csv(dateipfad)
    patient_name = df['Name'][0]
    date_time = df['Datum'][0]

    # --- Patient info with right-aligned tab ---
    p = doc.add_paragraph()
    ts = p.paragraph_format.tab_stops
    ts.add_tab_stop(section.page_width - section.left_margin - section.right_margin,
                    alignment=WD_TAB_ALIGNMENT.RIGHT)
    p.add_run(f"{patient_name}\t{date_time}")

    df['Asymmetry_dict'] = df['Asymmetry'].apply(ast.literal_eval)
    df['IA_LF'] = df['Asymmetry_dict'].apply(lambda x: x['Impulse']['Absprung']['L'])
    df['IA_RF'] = df['Asymmetry_dict'].apply(lambda x: x['Impulse']['Absprung']['R'])
    df['IL_LF'] = df['Asymmetry_dict'].apply(lambda x: x['Impulse']['Landing']['L'])
    df['IL_RF'] = df['Asymmetry_dict'].apply(lambda x: x['Impulse']['Landing']['R'])

    # --- Statistical calculation for take-off and landing ---
    _, _, diff_abs, m_abs, sd_abs = stat_impuls(df['IA_LF'].values, df['IA_RF'].values)
    _, _, diff_land, m_land, sd_land = stat_impuls(df['IL_LF'].values, df['IL_RF'].values)

    # Add percentage differences to DataFrame
    df['ΔIA(%)'] = diff_abs
    df['ΔIL(%)'] = diff_land

    # Ensure numeric columns
    cols = ['Sprunghöhe(Cm)', 'IA_RF(N.s)', 'IA_LF(N.s)',
            'ΔIA(%)', 'IL_RF(N.s)', 'IL_LF(N.s)', 'ΔIL(%)']
    df['Sprunghöhe(Cm)'] = df['IA_RF']
    df['IA_RF(N.s)'] = df['IA_RF']
    df['IA_LF(N.s)'] = df['IA_LF']
    df['IL_RF(N.s)'] = df['IL_RF']
    df['IL_LF(N.s)'] = df['IL_LF']
    df[cols] = df[cols].apply(pd.to_numeric, errors='coerce')

    # Restructure DataFrame for table
    pivot_df = df.set_index('Quelldatei')[cols].transpose()

    # Add mean ± SD column
    pivot_df['x̄ ± σ'] = [
        f"{df['Sprunghöhe(Cm)'].mean().round(2)} ± {df['Sprunghöhe(Cm)'].std().round(2)}",
        f"{df['IA_RF(N.s)'].mean().round(2)} ± {df['IA_RF(N.s)'].std().round(2)}",
        f"{df['IA_LF(N.s)'].mean().round(2)} ± {df['IA_LF(N.s)'].std().round(2)}",
        f"{m_abs} ± {sd_abs}",
        f"{df['IL_RF(N.s)'].mean().round(2)} ± {df['IL_RF(N.s)'].std().round(2)}",
        f"{df['IL_LF(N.s)'].mean().round(2)} ± {df['IL_LF(N.s)'].std().round(2)}",
        f"{m_land} ± {sd_land}"
    ]

    # --- Landscape section for the table ---
    section_land = doc.add_section()
    set_section_A4(section_land)
    section_land.orientation = WD_ORIENT.LANDSCAPE
    section_land.page_width, section_land.page_height = section.page_height, section.page_width
    section_land.different_first_page_header_footer = False

    # Header in landscape mode
    h2 = section_land.header.paragraphs[0]
    ts2 = h2.paragraph_format.tab_stops
    ts2.add_tab_stop(section_land.page_width - section_land.left_margin - section_land.right_margin,
                     alignment=WD_TAB_ALIGNMENT.RIGHT)
    h2.add_run(f"{patient_name}\t\t\t{date_time}")

    # --- Insert table ---
    create_responsive_styled_table(doc, pivot_df)

    # --- Save document ---
    filename = f"{patient_name} {date_time}.docx"
    doc.save(filename)
    print(f"Bericht erstellt: {filename}")

