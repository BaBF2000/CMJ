import json
from pathlib import Path

import pytest
from docx import Document
from docx.shared import RGBColor

from cmj_framework.export import word_report_helpers as helpers


def test_get_all_export_parameter_labels_returns_combined_labels():
    """
    Test that all exportable labels are returned in a single list.
    """
    labels = helpers.get_all_export_parameter_labels()

    expected = [label for label, _ in (helpers.PERF_METRICS + helpers.PEAKFORCE_METRICS)]
    assert labels == expected


def test_infer_patient_from_path_returns_patient_name():
    """
    Test that infer_patient_from_path extracts the patient name from the canonical path.
    """
    json_path = str(
        Path("CMJ_manager") / "Max Mustermann" / "01.01.2020" / "processed" / "file.json"
    )

    result = helpers.infer_patient_from_path(json_path)

    assert result == "Max Mustermann"


def test_infer_patient_from_path_returns_unknown_for_invalid_path():
    """
    Test that infer_patient_from_path falls back to Unknown for invalid paths.
    """
    result = helpers.infer_patient_from_path("file.json")

    assert result == "Unknown"


def test_safe_float_returns_float_for_valid_value():
    """
    Test that safe_float converts valid values to float.
    """
    assert helpers.safe_float("3.14") == 3.14
    assert helpers.safe_float(5) == 5.0


def test_safe_float_returns_default_for_invalid_value():
    """
    Test that safe_float returns the default value on failure.
    """
    assert helpers.safe_float("abc", default=99) == 99
    assert helpers.safe_float(None, default=7) == 7


def test_fetch_returns_nested_value():
    """
    Test that fetch returns a nested dictionary value.
    """
    data = {
        "A": {
            "B": {
                "C": 42
            }
        }
    }

    result = helpers.fetch(data, ["A", "B", "C"])

    assert result == 42


def test_fetch_returns_default_for_missing_path():
    """
    Test that fetch returns the default value when the path is missing.
    """
    data = {"A": {"B": 1}}

    result = helpers.fetch(data, ["A", "X"], default="missing")

    assert result == "missing"


def test_sanitize_filename_replaces_invalid_characters():
    """
    Test that sanitize_filename replaces forbidden characters with underscores.
    """
    result = helpers.sanitize_filename('bad:/\\*?"<>|name')

    assert result == "bad_________name"


def test_rgb_to_hex_returns_default_for_none():
    """
    Test that _rgb_to_hex returns white for None.
    """
    assert helpers._rgb_to_hex(None) == "FFFFFF"


def test_rgb_to_hex_converts_rgbcolor():
    """
    Test that _rgb_to_hex converts an RGBColor instance correctly.
    """
    result = helpers._rgb_to_hex(RGBColor(255, 170, 16))

    assert result == "FFAA10"


def test_rgb_to_hex_normalizes_hex_string():
    """
    Test that _rgb_to_hex normalizes hex-like strings.
    """
    assert helpers._rgb_to_hex("#a1b2c3") == "A1B2C3"
    assert helpers._rgb_to_hex("a1b2c3") == "A1B2C3"


def test_rgb_to_hex_returns_default_for_invalid_string():
    """
    Test that _rgb_to_hex returns white for invalid input.
    """
    assert helpers._rgb_to_hex("not-a-color") == "FFFFFF"


def test_ensure_editable_markdown_file_returns_existing_writable_copy(tmp_path, monkeypatch):
    """
    Test that ensure_editable_markdown_file returns the existing writable file if present.
    """
    writable_dir = tmp_path / "writable"
    writable_dir.mkdir()
    existing = writable_dir / "parameter.md"
    existing.write_text("existing content", encoding="utf-8")

    monkeypatch.setattr(helpers, "get_writable_export_resource_dir", lambda: writable_dir)

    result = helpers.ensure_editable_markdown_file("parameter.md")

    assert result == str(existing)
    assert existing.read_text(encoding="utf-8") == "existing content"


def test_ensure_editable_markdown_file_copies_bundled_file_when_missing(tmp_path, monkeypatch):
    """
    Test that ensure_editable_markdown_file copies the bundled file if no writable copy exists.
    """
    writable_dir = tmp_path / "writable"
    writable_dir.mkdir()

    bundled_dir = tmp_path / "bundled"
    bundled_dir.mkdir()
    bundled_file = bundled_dir / "parameter.md"
    bundled_file.write_text("bundled content", encoding="utf-8")

    monkeypatch.setattr(helpers, "get_writable_export_resource_dir", lambda: writable_dir)
    monkeypatch.setattr(helpers, "export_resource_file", lambda filename: bundled_file)

    result = helpers.ensure_editable_markdown_file("parameter.md")
    result_path = Path(result)

    assert result_path.exists()
    assert result_path.read_text(encoding="utf-8") == "bundled content"


def test_ensure_editable_markdown_file_creates_empty_file_when_bundle_missing(tmp_path, monkeypatch):
    """
    Test that ensure_editable_markdown_file creates an empty file if the bundled file is missing.
    """
    writable_dir = tmp_path / "writable"
    writable_dir.mkdir()

    missing_bundled_file = tmp_path / "bundled" / "missing.md"

    monkeypatch.setattr(helpers, "get_writable_export_resource_dir", lambda: writable_dir)
    monkeypatch.setattr(helpers, "export_resource_file", lambda filename: missing_bundled_file)

    result = helpers.ensure_editable_markdown_file("missing.md")
    result_path = Path(result)

    assert result_path.exists()
    assert result_path.read_text(encoding="utf-8") == ""


def test_insert_markdown_adds_headings_paragraphs_and_bullets(tmp_path):
    """
    Test that insert_markdown adds headings, body text, and bullet items to the document.
    """
    md_path = tmp_path / "example.md"
    md_path.write_text(
        "# Title\n\n"
        "Some **bold** text and *italic* text.\n\n"
        "- Item 1\n"
        "- Item 2\n\n"
        "## Subtitle\n\n"
        "Last paragraph.\n",
        encoding="utf-8",
    )

    doc = Document()
    helpers.insert_markdown(doc, str(md_path))

    texts = [p.text for p in doc.paragraphs]

    assert "Title" in texts
    assert "Some bold text and italic text." in texts
    assert "Item 1" in texts
    assert "Item 2" in texts
    assert "Subtitle" in texts
    assert "Last paragraph." in texts


def test_insert_markdown_preserves_basic_bold_and_italic_runs(tmp_path):
    """
    Test that insert_markdown creates bold and italic runs for simple inline markdown.
    """
    md_path = tmp_path / "inline.md"
    md_path.write_text(
        "This is **bold** and this is *italic*.\n",
        encoding="utf-8",
    )

    doc = Document()
    helpers.insert_markdown(doc, str(md_path))

    paragraph = doc.paragraphs[0]
    runs = paragraph.runs

    assert any(run.text == "bold" and run.bold for run in runs)
    assert any(run.text == "italic" and run.italic for run in runs)


def test_create_classic_table_with_stats_creates_stats_column():
    """
    Test that create_classic_table_with_stats adds a statistics column when requested.
    """
    pandas = pytest.importorskip("pandas")

    df = pandas.DataFrame(
        [
            ["Sprunghöhe", 10.0, 12.0, 14.0],
            ["RSI Modified", 1.0, 1.2, 1.4],
        ],
        columns=["Parameter", "Trial 1", "Trial 2", "Trial 3"],
    )

    doc = Document()
    table = helpers.create_classic_table_with_stats(doc, df, add_stats=True)

    assert len(table.rows) == 3
    assert len(table.columns) == 5
    assert table.cell(0, 4).text == "x̄ ± s"
    assert table.cell(1, 4).text == "12.0 ± 1.63"
    assert table.cell(2, 4).text == "1.2 ± 0.16"


def test_create_classic_table_with_stats_without_stats_column():
    """
    Test that create_classic_table_with_stats does not add a statistics column when disabled.
    """
    pandas = pytest.importorskip("pandas")

    df = pandas.DataFrame(
        [
            ["Sprunghöhe", 10.0, 12.0],
        ],
        columns=["Parameter", "Trial 1", "Trial 2"],
    )

    doc = Document()
    table = helpers.create_classic_table_with_stats(doc, df, add_stats=False)

    assert len(table.rows) == 2
    assert len(table.columns) == 3
    assert table.cell(0, 0).text == "Parameter"
    assert table.cell(1, 0).text == "Sprunghöhe"